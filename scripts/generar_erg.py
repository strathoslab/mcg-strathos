#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ERG — Escenarios de Riesgos Geopolíticos · Generador de dossier · Strathos Lab

Lee erg/<slug>.md (redactado por Fernando) y produce erg/ERG_<slug>.docx con la marca.
Soporta prosa y tablas de tubería (| a | b |) para canastas de exposición y señales.

Uso:  python scripts/generar_erg.py ormuz
"""

import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK   = RGBColor(0x0F, 0x1A, 0x2E)
CORE  = RGBColor(0xB8, 0x90, 0x2A)
OCRE  = RGBColor(0xB8, 0x90, 0x2A)
ORO_TXT = RGBColor(0x8A, 0x6B, 0x18)   # dorado oscurecido: legible en texto chico
GRIS  = RGBColor(0x4A, 0x55, 0x68)
TINTA = RGBColor(0x1A, 0x1A, 0x1A)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
FUENTE = "Calibri"
RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def sf(run, size=11, color=TINTA, bold=False, italic=False, name=FUENTE):
    run.font.name = name; run.font.size = Pt(size)
    run.font.color.rgb = color; run.font.bold = bold; run.font.italic = italic


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "0F1A2E")):
        b.set(qn(k), v)
    pb.append(b); pPr.append(pb)
    p.paragraph_format.space_after = Pt(2)


def render_tabla(doc, filas):
    """filas: lista de listas de celdas (la primera es encabezado)."""
    ncol = len(filas[0])
    t = doc.add_table(rows=1, cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, txt in enumerate(filas[0]):
        c = t.rows[0].cells[j]; shade(c, "0F1A2E")
        p = c.paragraphs[0]; r = p.add_run(txt); sf(r, size=9, color=BLANCO, bold=True)
    for fila in filas[1:]:
        cells = t.add_row().cells
        for j, txt in enumerate(fila):
            p = cells[j].paragraphs[0]
            # coloreado de dirección
            color = TINTA; bold = False
            if txt.strip() in ("↑", "sube", "Sube"):
                color = RGBColor(0x1F, 0x6F, 0x4A); bold = True
            elif txt.strip() in ("↓", "baja", "Baja"):
                color = RGBColor(0xA3, 0x3A, 0x2E); bold = True
            elif txt.strip().lower() in ("vol", "volátil", "volatilidad"):
                color = ORO_TXT; bold = True
            r = p.add_run(txt); sf(r, size=9, color=color, bold=(bold or j == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def es_fila_tabla(l):
    return l.strip().startswith("|") and l.strip().endswith("|")


def celdas(l):
    return [c.strip() for c in l.strip().strip("|").split("|")]


def main():
    slug = sys.argv[1] if len(sys.argv) > 1 else "ormuz"
    src = os.path.join(RAIZ, "erg", f"{slug}.md")
    if not os.path.exists(src):
        print(f"ERROR: no existe {src}", file=sys.stderr); sys.exit(1)
    lineas = open(src, encoding="utf-8").read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = FUENTE
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

    # cabecera
    p = doc.add_paragraph(); r = p.add_run("STRATHOS LAB"); sf(r, 11, ORO_TXT, bold=True)
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(); r = p.add_run("Escenarios de Riesgos Geopolíticos"); sf(r, 21, INK, bold=True)
    p.paragraph_format.space_after = Pt(0)

    primer_bloque = True
    i = 0
    n = len(lineas)
    while i < n:
        l = lineas[i]
        if l.startswith("# "):
            hr(doc) if primer_bloque else None
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(l[2:].strip()); sf(r, 15, INK, bold=True)
            primer_bloque = False
            i += 1; continue
        if l.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
            r = p.add_run(l[3:].strip()); sf(r, 12, INK, bold=True)
            i += 1; continue
        if es_fila_tabla(l):
            # juntar el bloque de tabla
            bloque = []
            while i < n and es_fila_tabla(lineas[i]):
                if not set(lineas[i].replace("|", "").strip()) <= set("-: "):  # saltar separador
                    bloque.append(celdas(lineas[i]))
                i += 1
            if bloque:
                render_tabla(doc, bloque)
            continue
        if l.strip() == "":
            i += 1; continue
        # párrafo de prosa (juntar líneas hasta blanco / encabezado / tabla)
        buff = []
        while i < n and lineas[i].strip() != "" and not lineas[i].startswith("#") and not es_fila_tabla(lineas[i]):
            buff.append(lineas[i].strip()); i += 1
        texto = " ".join(buff)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.28
        if texto.startswith("**") and texto.endswith("**"):
            r = p.add_run(texto.strip("*")); sf(r, 12.5, INK, bold=True)
            p.paragraph_format.space_after = Pt(12)
        else:
            r = p.add_run(texto); sf(r, 11, TINTA)

    # firma
    doc.add_paragraph()
    hr(doc)
    p = doc.add_paragraph(); r = p.add_run("Fernando Legrand — Strathos Lab"); sf(r, 10.5, INK, bold=True)
    p = doc.add_paragraph(); r = p.add_run("strathoslab.com · Escenarios de Riesgos Geopolíticos (ERG)"); sf(r, 9, GRIS)

    os.makedirs(os.path.join(RAIZ, "erg"), exist_ok=True)
    out = os.path.join(RAIZ, "erg", f"ERG_{slug}.docx")
    doc.save(out)
    print("Dossier generado:", out)


if __name__ == "__main__":
    main()
