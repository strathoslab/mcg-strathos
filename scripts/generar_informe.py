#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MCG — Generador del Informe Mensual · Strathos Lab

Combina dos insumos y produce un DOCX con la marca puesta:
  1. docs/mcg_estado.json         -> los datos (señales, precios). LA MÁQUINA.
  2. informe/analisis_AAAA-MM.md  -> la lectura de Fernando.        STRATHOS.

Uso:
    python scripts/generar_informe.py 2026-07
Si no se pasa el mes, usa el mes en curso.

Filosofía: el generador nunca inventa análisis. Toma los números del JSON y la
interpretación del .md. Si falta una sección del .md, deja un marcador visible
para que no se publique un informe incompleto por error.
"""

import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- identidad visual Strathos (sobria) ----------
INK   = RGBColor(0x1F, 0x2A, 0x33)   # slate — títulos
CORE  = RGBColor(0x26, 0x41, 0x3C)   # verde profundo — acentos
OCRE  = RGBColor(0x7A, 0x6A, 0x3F)   # ocre — señales
GRIS  = RGBColor(0x5A, 0x64, 0x70)   # gris — metadatos
SUBE  = RGBColor(0x1F, 0x6F, 0x4A)
BAJA  = RGBColor(0xA3, 0x3A, 0x2E)
TINTA = RGBColor(0x22, 0x2A, 0x30)
FUENTE = "Calibri"

MESES = ["", "enero", "febrero", "marzo", "abril", "mayo", "junio",
         "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# ---------------------------------------------------------------- utilidades
def set_font(run, size=11, color=TINTA, bold=False, italic=False, name=FUENTE):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def pct_run(paragraph, value):
    """Agrega un run con el porcentaje coloreado según signo."""
    if value is None:
        r = paragraph.add_run("—")
        set_font(r, size=9.5, color=GRIS)
        return
    r = paragraph.add_run(f"{'+' if value >= 0 else ''}{value:.2f}%")
    set_font(r, size=9.5, color=(SUBE if value >= 0 else BAJA), bold=True)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F2A33")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------- parser .md
def parse_analisis(texto):
    """
    Divide el .md en secciones por encabezados '# '. Devuelve un dict
    {titulo_normalizado: {'titulo': str, 'subtitulo': str|None, 'parrafos': [str]}}
    Un '## ' dentro de una sección se toma como subtítulo sustantivo.
    """
    secciones = {}
    actual = None
    for linea in texto.splitlines():
        if linea.startswith("# "):
            titulo = linea[2:].strip()
            actual = {"titulo": titulo, "subtitulo": None, "parrafos": []}
            secciones[titulo.lower()] = actual
        elif linea.startswith("## ") and actual is not None:
            actual["subtitulo"] = linea[3:].strip()
        elif actual is not None:
            if linea.strip() == "":
                actual["parrafos"].append("")
            else:
                actual["parrafos"].append(linea.rstrip())
    # unir líneas en párrafos (bloques separados por líneas vacías)
    for s in secciones.values():
        bloques, buff = [], []
        for l in s["parrafos"]:
            if l == "":
                if buff:
                    bloques.append(" ".join(buff)); buff = []
            else:
                buff.append(l)
        if buff:
            bloques.append(" ".join(buff))
        s["parrafos"] = [b for b in bloques if b.strip()]
    return secciones


def escribir_prosa(doc, parrafos, bajada=False):
    """Escribe párrafos de cuerpo. El primero puede ir como bajada (standfirst)."""
    if not parrafos:
        p = doc.add_paragraph()
        r = p.add_run("[Falta la redacción de esta sección en el archivo de análisis.]")
        set_font(r, size=11, color=BAJA, italic=True)
        return
    for i, texto in enumerate(parrafos):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.28
        if bajada and i == 0:
            r = p.add_run(texto)
            set_font(r, size=12.5, color=INK, bold=True)
            p.paragraph_format.space_after = Pt(12)
        else:
            r = p.add_run(texto)
            set_font(r, size=11, color=TINTA)


# ---------------------------------------------------------------- bloques doc
def encabezado_seccion(doc, kicker, titulo):
    """Kicker (teatro) + subtítulo sustantivo."""
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(kicker.upper())
        set_font(r, size=9, color=CORE, bold=True)
        r.font.name = FUENTE
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(titulo)
    set_font(r, size=15, color=INK, bold=True)


def tabla_senales(doc, senales):
    if not senales:
        p = doc.add_paragraph()
        r = p.add_run("Sin movimientos fuera de rango habitual este período.")
        set_font(r, size=10.5, color=GRIS, italic=True)
        return
    cols = ["Activo", "Teatro / Conflicto", "1 día", "30 días", "Desvíos (z)"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]
        shade(cell, "1F2A33")
        pr = cell.paragraphs[0]; pr.paragraph_format.space_after = Pt(2)
        r = pr.add_run(c); set_font(r, size=9, color=RGBColor(0xFF,0xFF,0xFF), bold=True)
    for s in senales:
        row = t.add_row().cells
        # activo
        p0 = row[0].paragraphs[0]
        r = p0.add_run(s["ticker"]); set_font(r, size=9.5, color=INK, bold=True)
        r2 = p0.add_run(f"\n{s['nombre']}"); set_font(r2, size=8, color=GRIS)
        # teatro/conflicto
        p1 = row[1].paragraphs[0]
        r = p1.add_run(s.get("teatro", "")); set_font(r, size=9, color=CORE)
        r2 = p1.add_run(f"\n{s['conflicto']}"); set_font(r2, size=8.5, color=GRIS)
        # 1d, 30d
        for idx, key in ((2, "var_1d_pct"), (3, "var_30d_pct")):
            pc = row[idx].paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pct_run(pc, s.get(key))
        # z
        pz = row[4].paragraphs[0]; pz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rz = pz.add_run(f"{s['z_score']:+.2f}")
        set_font(rz, size=9.5, color=OCRE, bold=True)


def tabla_teatro(doc, activos_mov):
    """Tabla compacta de los activos de un teatro que se movieron fuera de rango."""
    if not activos_mov:
        return
    cols = ["Activo", "Recurso", "Precio", "1 d", "30 d", "Vol 30d"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, c in enumerate(cols):
        cell = t.rows[0].cells[j]; shade(cell, "3A464E")
        pr = cell.paragraphs[0]; r = pr.add_run(c)
        set_font(r, size=8.5, color=RGBColor(0xFF,0xFF,0xFF), bold=True)
    for a in activos_mov:
        d = a["datos"]
        row = t.add_row().cells
        r = row[0].paragraphs[0].add_run(a["ticker"]); set_font(r, size=9, color=INK, bold=True)
        r = row[1].paragraphs[0].add_run(a["recurso"]); set_font(r, size=8.5, color=TINTA)
        pp = row[2].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = pp.add_run(f"{d['precio']}"); set_font(r, size=9, color=TINTA)
        for idx, key in ((3, "var_1d_pct"), (4, "var_30d_pct")):
            pc = row[idx].paragraphs[0]; pc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pct_run(pc, d.get(key))
        pv = row[5].paragraphs[0]; pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        v = d.get("volatilidad_30d_anualizada_pct")
        r = pv.add_run(f"{v}%" if v is not None else "—"); set_font(r, size=9, color=GRIS)


# ---------------------------------------------------------------- principal
def main():
    mes_arg = sys.argv[1] if len(sys.argv) > 1 else datetime.now().strftime("%Y-%m")
    try:
        anio, mes = [int(x) for x in mes_arg.split("-")]
    except ValueError:
        print(f"ERROR: mes inválido '{mes_arg}'. Formato AAAA-MM.", file=sys.stderr)
        sys.exit(1)

    estado_path = os.path.join(RAIZ, "docs", "mcg_estado.json")
    analisis_path = os.path.join(RAIZ, "informe", f"analisis_{mes_arg}.md")

    if not os.path.exists(estado_path):
        print("ERROR: falta docs/mcg_estado.json. Corré primero el job de precios.", file=sys.stderr)
        sys.exit(1)
    with open(estado_path, encoding="utf-8") as f:
        estado = json.load(f)

    if os.path.exists(analisis_path):
        with open(analisis_path, encoding="utf-8") as f:
            secciones = parse_analisis(f.read())
    else:
        print(f"AVISO: no existe {analisis_path}. Se genera con marcadores de redacción.")
        secciones = {}

    umbral = estado["meta"].get("umbral_senal_sigma", 1.5)

    doc = Document()
    doc.styles["Normal"].font.name = FUENTE
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

    # ---- Portada / cabecera ----
    p = doc.add_paragraph()
    r = p.add_run("STRATHOS LAB")
    set_font(r, size=11, color=CORE, bold=True)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Monitor de Conflictos Geopolíticos")
    set_font(r, size=22, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    r = p.add_run(f"Informe mensual · {MESES[mes]} de {anio}")
    set_font(r, size=13, color=GRIS)
    p.paragraph_format.space_after = Pt(4)
    add_hr(doc)

    p = doc.add_paragraph()
    r = p.add_run(f"Datos al {estado['meta']['actualizado_utc']} · "
                  f"{estado['meta']['vinculos']} vínculos activo-conflicto · "
                  f"fuente de precios: {estado['meta']['fuente_precios']}")
    set_font(r, size=8.5, color=GRIS, italic=True)
    p.paragraph_format.space_after = Pt(14)

    # ---- Lectura del mes ----
    sec = secciones.get("lectura del mes")
    encabezado_seccion(doc, None, sec["subtitulo"] if sec and sec["subtitulo"] else "Lectura del mes")
    escribir_prosa(doc, sec["parrafos"] if sec else [], bajada=True)

    # ---- Tablero de señales (máquina) ----
    encabezado_seccion(doc, "La máquina", "Tablero de señales")
    p = doc.add_paragraph()
    r = p.add_run(f"Activos que se movieron más de {umbral} desvíos respecto de su volatilidad "
                  f"reciente. Registro de exposición, no de causalidad: la señal indica que el activo "
                  f"se movió fuera de su rango habitual, no que el conflicto lo haya causado.")
    set_font(r, size=9.5, color=GRIS, italic=True)
    p.paragraph_format.space_after = Pt(8)
    tabla_senales(doc, estado.get("senales", []))

    # ---- Teatros ----
    # tickers con señal, para resaltar por teatro
    con_senal = {s["ticker"] for s in estado.get("senales", [])}
    orden_teatros = ["Rusia–Ucrania", "Complejo Israel–Irán", "Indo-Pacífico",
                     "Américas", "África de Recursos Críticos"]
    teatros_json = {t["teatro"]: t for t in estado.get("teatros", [])}

    for nombre in orden_teatros:
        tj = teatros_json.get(nombre)
        sec = secciones.get(f"teatro: {nombre.lower()}") or secciones.get(nombre.lower())
        subtitulo = sec["subtitulo"] if sec and sec["subtitulo"] else nombre
        encabezado_seccion(doc, nombre, subtitulo)

        # activos del teatro que movieron fuera de rango
        activos_mov = []
        if tj:
            for c in tj["conflictos"]:
                for a in c["activos"]:
                    if a["ticker"] in con_senal and a.get("datos"):
                        activos_mov.append(a)
        tabla_teatro(doc, activos_mov)
        if activos_mov:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        escribir_prosa(doc, sec["parrafos"] if sec else [])

    # ---- Señales a vigilar ----
    sec = secciones.get("señales a vigilar") or secciones.get("senales a vigilar")
    encabezado_seccion(doc, "Prospectiva", "Señales a vigilar")
    escribir_prosa(doc, sec["parrafos"] if sec else [])

    # ---- Encuadre metodológico (boilerplate fijo) ----
    add_hr(doc)
    p = doc.add_paragraph()
    r = p.add_run("Encuadre metodológico. ")
    set_font(r, size=9, color=INK, bold=True)
    r = p.add_run(
        "El MCG es un monitor de exposición. Mapea conflictos activos a activos cotizables en "
        "minería, banca y energía, y registra sus movimientos. No establece relaciones causales "
        "entre un conflicto y un movimiento de precio: esa lectura es analítica y está firmada. "
        "Los precios provienen de Yahoo Finance; la interpretación se apoya en fuentes primarias "
        "verificadas. Los movimientos marcados como señal son estadísticamente inusuales para cada "
        "activo, no necesariamente relevantes.")
    set_font(r, size=9, color=GRIS)
    p.paragraph_format.space_after = Pt(10)

    # ---- Firma ----
    p = doc.add_paragraph()
    r = p.add_run("Fernando Legrand — Strathos Lab")
    set_font(r, size=10.5, color=INK, bold=True)
    p = doc.add_paragraph()
    r = p.add_run(f"strathoslab.com · {MESES[mes]} de {anio}")
    set_font(r, size=9, color=GRIS)

    os.makedirs(os.path.join(RAIZ, "informe"), exist_ok=True)
    salida = os.path.join(RAIZ, "informe", f"MCG_{mes_arg}.docx")
    doc.save(salida)
    print(f"Informe generado: {salida}")
    faltantes = [k for k in ["lectura del mes", "señales a vigilar"] if k not in secciones]
    if faltantes or not os.path.exists(analisis_path):
        print("AVISO: hay secciones sin redactar. Revisá los marcadores en rojo antes de publicar.")


if __name__ == "__main__":
    main()
